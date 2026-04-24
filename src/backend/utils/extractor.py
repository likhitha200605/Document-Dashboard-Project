import os
import pdfplumber
import docx


SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt"}


def extract_text(file_path: str, filename: str) -> tuple[str, int]:
    """
    Detect file type and extract text. Returns (text, page_count).
    page_count is 0 for non-PDF formats.
    Raises RuntimeError on failure.
    """
    ext = os.path.splitext(filename.lower())[1]

    if ext not in SUPPORTED_EXTENSIONS:
        raise RuntimeError(
            f"Unsupported file type '{ext}'. Supported: PDF, DOCX, DOC, TXT."
        )

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    elif ext == ".txt":
        return _extract_txt(file_path)


def _extract_pdf(file_path: str) -> tuple[str, int]:
    text_parts = []
    try:
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            if page_count == 0:
                raise ValueError("PDF has no pages")
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        raise RuntimeError(f"Failed to read PDF: {str(e)}")
    return "\n".join(text_parts), page_count


def _extract_docx(file_path: str) -> tuple[str, int]:
    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        text = "\n".join(paragraphs)
    except Exception as e:
        raise RuntimeError(f"Failed to read DOCX: {str(e)}")
    return text, 0


def _extract_txt(file_path: str) -> tuple[str, int]:
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
    except Exception as e:
        raise RuntimeError(f"Failed to read TXT: {str(e)}")
    return text, 0


# Keep old function name as alias for backward compatibility
def extract_text_from_pdf(file_path: str) -> tuple[str, int]:
    return _extract_pdf(file_path)


def get_page_count(file_path: str) -> int:
    try:
        with pdfplumber.open(file_path) as pdf:
            return len(pdf.pages)
    except Exception:
        return 0

